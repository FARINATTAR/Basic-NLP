import docx  # For Word document handling
import nltk  # For text processing
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk import pos_tag


def extract_pos_from_docx(input_file, output_file):
    # Loading stop words
    stop_words = set(stopwords.words('english'))

    # Opening the Word document
    doc = docx.Document(input_file)
    text = ""

    # Extracting text from paragraphs
    for para in doc.paragraphs:
        text += para.text + " "

    # Tokenize words and remove stop words
    words = word_tokenize(text)
    filtered_words = [word for word in words if word.lower() not in stop_words]

    # POS tagging (Parts of Speech)
    tagged_words = pos_tag(filtered_words)

    # Extract nouns, verbs, and adjectives
    nouns = [word for word, pos in tagged_words if pos.startswith('NN')]
    verbs = [word for word, pos in tagged_words if pos.startswith('VB')]
    adjectives = [word for word, pos in tagged_words if pos.startswith('JJ')]

    # Create a new Word document
    new_doc = docx.Document()
    new_doc.add_heading("Extracted Words", level=1)

    # Add nouns
    new_doc.add_heading("Nouns:", level=2)
    new_doc.add_paragraph(", ".join(nouns))

    # Add verbs
    new_doc.add_heading("Verbs:", level=2)
    new_doc.add_paragraph(", ".join(verbs))

    # Add adjectives
    new_doc.add_heading("Adjectives:", level=2)
    new_doc.add_paragraph(", ".join(adjectives))

    # Save the new Word file
    new_doc.save(output_file)
    print(f"Extraction complete! File saved as '{output_file}'.")

# Example usage
input_file = "input.docx"   # Input Word file path
output_file = "extracted_words.docx"  # Output Word file path
extract_pos_from_docx(input_file, output_file)
